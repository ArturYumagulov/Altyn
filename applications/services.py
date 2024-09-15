import os
import django

from django.conf import settings
import datetime
import io

from django.core.mail import EmailMessage
from django.http import HttpResponse
from docx import Document

# Устанавливаем переменную окружения для настройки Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'altyn.settings')

# Инициализируем Django
django.setup()

# Теперь можно импортировать модели и использовать настройки


# Используем модели и настройки
contract_ooo_path = os.path.join(settings.BASE_DIR, 'contract/contract_ooo.docx')


def send_word_via_email(request):
    # Открываем шаблон документа
    template_path = contract_ooo_path
    doc = Document(template_path)

    # Данные для замены в шаблоне
    context = {
        'имя': 'Иван Иванов',
        'должность': 'Программист',
        'дата': '13 сентября 2024'
    }

    # Функция для замены маркеров в тексте абзацев
    def replace_text_in_paragraph(paragraph, context):
        for key, value in context.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{key}}}', value)

    # Проходим по всем абзацам документа и заменяем маркеры
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, context)

    # Создаем объект in-memory для хранения документа
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Создаем письмо
    email = EmailMessage(
        subject='Ваш документ Word',
        body='Пожалуйста, найдите вложение с вашим документом.',
        from_email=settings.EMAIL_HOST_USER,  # Ваш email
        to=['zico.13288@gmail.com', 'vladakrylova@ya.ru']  # Получатель письма
    )

    # Добавляем документ как вложение
    email.attach('договор_на_фильм.docx', file_stream.getvalue(),
                 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # Отправляем письмо
    if email.send():
        return True
    return False
#
#
# # Открываем шаблон
# doc = Document('/Users/arturumagulov/PycharmProjects/Altyn/contract_ooo.docx')
#
# # Словарь с данными для замены
# context = {
#     'create_date': f'{datetime.datetime.now().year}-{datetime.datetime.now().month}-{datetime.datetime.now().day}',
# }
#
#
# bold_context = {
#     'full_name': "Юмагулов Артур Ильгизович"
# }
#
#
# # Функция для замены текста в документе
# def replace_text(paragraph, context):
#     for key, value in context.items():
#         if key in paragraph.text:
#             paragraph.text = paragraph.text.replace(f'{{{key}}}', value)
#
#
# def replace_bold_text(paragraph, context):
#     for key, value in context.items():
#         if f'{{{key}}}' in paragraph.text:
#             # Разбиваем текст абзаца на части до и после маркера
#             parts = paragraph.text.split(f'{{{key}}}')
#
#             # Очищаем текущий текст абзаца
#             paragraph.clear()
#
#             # Добавляем текст до маркера
#             if parts[0]:
#                 paragraph.add_run(parts[0])
#
#             # Добавляем значение переменной с жирным шрифтом
#             run = paragraph.add_run(value)
#             run.bold = True
#
#             # Добавляем текст после маркера
#             if parts[1]:
#                 paragraph.add_run(parts[1])
#
#
# # Проходим по всем абзацам документа и заменяем маркеры
# for paragraph in doc.paragraphs:
#     replace_text(paragraph, context)
#     replace_bold_text(paragraph, bold_context)
#
# # Сохранение нового документа
# doc.save('/Users/arturumagulov/PycharmProjects/Altyn/generated_document.docx')


